# encoding: utf-8

"""
Step implementations for table-related features
"""

from __future__ import absolute_import, print_function, unicode_literals

from behave import given, then

from docx import Document
from docx.table import _Cell, _CellCollection, _Row

from .helpers import test_docx


# given ===================================================

@given('a table having two rows')
def given_a_table_having_two_rows(context):
    docx_path = test_docx('blk-containing-table')
    document = Document(docx_path)
    context.table_ = document.body.tables[0]


@given('a table row having two cells')
def given_a_table_row_having_two_cells(context):
    docx_path = test_docx('blk-containing-table')
    document = Document(docx_path)
    context.row = document.body.tables[0].rows[0]


# then =====================================================

@then('I can access a collection cell by index')
def then_can_access_collection_cell_by_index(context):
    cells = context.row.cells
    for idx in range(2):
        cell = cells[idx]
        assert isinstance(cell, _Cell)


@then('I can access the cell collection of the row')
def then_can_access_cell_collection_of_row(context):
    row = context.row
    cells = row.cells
    assert isinstance(cells, _CellCollection)


@then('I can access the rows by index')
def then_can_access_rows_by_index(context):
    rows = context.table_.rows
    row_count = len(rows)
    for idx in range(row_count):
        row = rows[idx]
        assert isinstance(row, _Row)


@then('I can get the length of the cell collection')
def then_can_get_length_of_cell_collection(context):
    row = context.row
    cells = row.cells
    assert len(cells) == 2


@then('I can iterate over the cell collection')
def then_can_iterate_over_cell_collection(context):
    row = context.row
    actual_count = 0
    for cell in row.cells:
        actual_count += 1
        assert isinstance(cell, _Cell)
    assert actual_count == 2


@then('the length of its row collection is 2')
def then_len_of_row_collection_is_2(context):
    rows = context.table_.rows
    assert len(rows) == 2


@then('each item in its row collection is a table row')
def then_each_item_in_row_collection_is_a_table_row(context):
    rows = context.table_.rows
    count = 0
    for item in rows:
        count += 1
        assert isinstance(item, _Row)
    assert count == 2
